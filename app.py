import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from openai import OpenAI
client = OpenAI(
  base_url="https://openrouter.ai/api/v1",
  api_key=st.secrets["OPENROUTER_API_KEY"],
  default_headers={
    "HTTP-Referer": "https://report-generator.streamlit.app", # Твой адрес
    "X-Title": "Report Generator"
  }
)
GEMINI_MODEL = "anthropic/claude-3.5-sonnet"
import io
import json
import re

if "reset_counter" not in st.session_state:
    st.session_state.reset_counter = 0

# --- 1. ФУНКЦИИ ПАРСИНГА (ТВОИ ОРИГИНАЛЬНЫЕ) ---

def get_contract_start_text(file):
    doc = Document(file)
    full_text = []
    for table in doc.tables:
        for row in table.rows:
            full_text.append(" ".join(cell.text.strip() for cell in row.cells))
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            if re.match(r"^2\.", txt): 
                break
            full_text.append(txt)
    return "\n".join(full_text)[:2000]

def get_text_from_file(file):
    doc = Document(file)
    content = []
    for p in doc.paragraphs:
        if p.text.strip(): content.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            content.append(" ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(content)

def format_fio_short(fio_str):
    if not fio_str: return "___________"
    parts = fio_str.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    return fio_str

def clean_markdown(text):
    return text.replace('**', '').replace('###', '').replace('##', '').replace('|', '').strip()

# --- 2. УМНАЯ ГЕНЕРАЦИЯ (ЛОГИКА ВНУТРИ) ---

def smart_generate_step_strict(section_text, requirements_text):
    system_prompt = f"""Ты - юридический редактор. Перепиши пункты ТЗ в Отчет.
    ПРАВИЛА:
    1. ВРЕМЯ: У заголовком - настоящее, у текста - СТРОГО ПРОШЕДШЕЕ ('организовано', 'оказано', 'размещено').
    2. ЗАПРЕТ: Слова 'должен', 'обязан', 'будет', 'необходимо' КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНЫ.
    3. ТОЧНОСТЬ: Перенеси ВСЕ цифры, площади, сроки и названия без изменений.
    4. ПУНКТУАЦИЯ: Соблюдай правила русского языка. Не обрывай предложения.
    ТРЕБОВАНИЯ К ДОКУМЕНТАМ: {requirements_text}"""

    full_prompt = f"{system_prompt}\n\nТРАНСФОРМИРУЙ ЭТОТ КУСОК ТЗ В ОТЧЕТ:\n{section_text}"

    # Шаг 1: Генерация
    res = client.chat.completions.create(
      model=GEMINI_MODEL,
      messages=[{"role": "user", "content": full_prompt}]
    )
    draft = res.choices[0].message.content

    # Шаг 2: ЖЕСТКИЙ КОНТРОЛЬ
    v_prompt = f"Сравни ТЗ и Отчет. Если есть ошибки, напиши их. Если всё ок, пиши 'ОШИБОК: 0'.\nТЗ: {section_text}\nОТЧЕТ: {draft}"
    v_res = client.chat.completions.create(
        model=GEMINI_MODEL,
        messages=[{"role": "user", "content": v_prompt}]
    )
    v_text = v_res.choices[0].message.content
    
    # Шаг 3: Исправление (если инспектор нашел брак)
    if "ОШИБОК: 0" not in v_text:
        fix_prompt = f"{system_prompt}\nИСПРАВЬ ОШИБКИ: {v_text}\nТЗ: {section_text}\nЧЕРНОВИК: {draft}"
        fix = client.chat.completions.create(
            model=GEMINI_MODEL,
            messages=[{"role": "user", "content": fix_prompt}]
        )
        return fix.choices[0].message.content
    return draft

# --- 3. СБОРКА ДОКУМЕНТА (ТВОЕ ОФОРМЛЕНИЕ) ---

def build_title_page(t):
    doc = Document()
    
    # Настройка узких полей (чтобы точно влезло)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)    # 1.27 см
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(72)   # 2.54 см
        section.right_margin = Pt(36)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Сбор данных [cite: 151]
    contract_no = t.get('contract_no', '___')
    contract_date = t.get('contract_date', '___')
    ikz = t.get('ikz', '___________')

    # Шапка 
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Информационно-аналитический отчет об исполнении условий\n").bold = True
    p.add_run(f"Контракта № {contract_no} от «{contract_date}» 2025 г.\n").bold = True
    p.add_run(f"Идентификационный код закупки: {ikz}").bold = True

    # Уменьшенный отступ перед ТОМ I
    for _ in range(3): doc.add_paragraph() 
    doc.add_paragraph("ТОМ I").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Предмет, Заказчик, Исполнитель [cite: 3-8, 153-158]
    for label, val in [
        ("Наименование предмета КОНТРАКТА:", t.get('project_name')),
        ("Заказчик:", t.get('customer')),
        ("Исполнитель:", t.get('company'))
    ]:
        p_l = doc.add_paragraph()
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.add_run(f"\n{label}").bold = True
        
        p_v = doc.add_paragraph()
        p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ограничиваем длину текста предмета, если он слишком длинный
        p_v.add_run(str(val)).italic = True

    # Динамический отступ перед подписями (уменьшен до 4)
    for _ in range(4): doc.add_paragraph()
    
    # Таблица подписей [cite: 9, 159]
    tab = doc.add_table(rows=2, cols=2)
    tab.autofit = True

    cust_post = str(t.get('customer_post', 'Заказчик')).capitalize()
    exec_post = str(t.get('director_post', 'Исполнитель')).capitalize()
    cust_fio = format_fio_short(t.get('customer_fio'))
    exec_fio = format_fio_short(t.get('director'))

    # Левая ячейка (Заказчик)
    p1 = tab.rows[0].cells[0].paragraphs[0]
    p1.add_run(f"Отчет принят Заказчиком\n{cust_post}\n\n___________ / {cust_fio}")
    
    # Правая ячейка (Исполнитель)
    p2 = tab.rows[0].cells[1].paragraphs[0]
    p2.add_run(f"Отчет передан Исполнителем\n{exec_post}\n\n___________ / {exec_fio}")

    # М.П. (нижняя строка таблицы)
    tab.rows[1].cells[0].text = "м.п."
    tab.rows[1].cells[1].text = "м.п."

    return doc
    
def apply_yellow_highlight(doc):
    keywords = ["Акт", "Фотоотчет", "Ведомость", "Скриншот", "Смета", "Резюме", "USB", "Флеш-накопитель"]
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for word in keywords:
                if word.lower() in run.text.lower():
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def create_final_report(t, report_body, req_body):
    doc = build_title_page(t)
    doc.add_page_break()
    
    # Чтобы не было "Отчет об оказании услуг по услуг"
    p_name = t.get('project_name', '')
    if isinstance(p_name, dict): p_name = p_name.get('name', '')
    p_name = str(p_name).strip() if p_name else "услугам"

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.add_run(f"Отчет об оказании услуг по {p_name}").bold = True
    
    # Очищаем основной текст от дублей и пустых строк
    lines = clean_markdown(report_body).split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        para = doc.add_paragraph()
        run = para.add_run(line)
        if re.match(r"^\d+\.", line): run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    if req_body:
        doc.add_page_break()
        p_req = doc.add_paragraph()
        p_req.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_req.add_run('ТРЕБОВАНИЯ К ПРЕДОСТАВЛЯЕМОЙ ДОКУМЕНТАЦИИ').bold = True
        doc.add_paragraph(clean_markdown(req_body)).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    apply_yellow_highlight(doc)
    return doc

# --- 4. ИНТЕРФЕЙС (ВОЗВРАТ К ТВОЕЙ СТРУКТУРЕ) ---

st.set_page_config(page_title="Генератор Отчетов 3.0", layout="wide")

with st.sidebar:
    st.title("Авторизация")
    if "auth" not in st.session_state: st.session_state.auth = False
    pwd = st.text_input("Пароль", type="password")
    if pwd == st.secrets["APP_PASSWORD"]: st.session_state.auth = True
    if not st.session_state.auth: st.stop()
    if st.button("♻️ СБРОСИТЬ ВСЕ ДАННЫЕ", use_container_width=True, type="primary"):
        # 1. Полная очистка session_state
        for key in list(st.session_state.keys()):
            if key != "reset_counter":
                del st.session_state[key]
        
        # 2. Явное обнуление переменных кэша текста (чтобы ИИ не подтянул старое)
        st.session_state.raw_tz_source = ""
        st.session_state.raw_report_body = ""
        st.session_state.raw_requirements = ""
        st.session_state.t_info = {}

        # Это принудительно очистит text_area в колонках
        st.session_state[f"t_area_{st.session_state.reset_counter}"] = ""
        st.session_state[f"tz_area_{st.session_state.reset_counter}"] = ""
        
        # 3. Смена ключей виджетов (то, что мы делали со счетчиком)
        st.session_state.reset_counter += 1
        
        # 4. Очистка кэша самого Streamlit (на всякий случай)
        st.cache_data.clear()
        
        st.rerun()
    
col1, col2, col3 = st.columns(3)

# КОЛОНКА 1: ТИТУЛЬНИК
with col1:
    st.header("📄 1. Титульный лист")
    t_tab1, t_tab2 = st.tabs(["📁 Файл", "⌨️ Текст"])
    
    t_context = ""
    with t_tab1:
        f_title = st.file_uploader("Контракт (DOCX)", type="docx", key="u_title")
        t_context = "" # Инициализируем пустой строкой
        if f_title: 
            t_context = get_contract_start_text(f_title)
    
    with t_tab2:
        # 1. Сначала определяем ключ текущего виджета
        area_key = f"t_area_{st.session_state.reset_counter}"
        
        # 2. Отрисовываем виджет
        m_title = st.text_area(
            "Вставьте начало контракта:", 
            value=st.session_state.get(area_key, ""), 
            height=150, 
            key=area_key
        )
    
    # 3. Если в поле что-то вписали, обновляем контекст
    if m_title: 
        t_context = m_title

    if st.button("🔍 Извлечь реквизиты", use_container_width=True):
        if t_context:
            with st.spinner("Ищем данные..."):
                res = client.chat.completions.create(
                    model=GEMINI_MODEL,
                    messages=[{
                        "role": "user", 
                        "content": f"""Извлеки данные СТРОГО в формате JSON с этими ключами:
                        'contract_no' (номер), 
                        'contract_date' (дата), 
                        'ikz' (ИКЗ), 
                        'project_name' (предмет контракта), 
                        'customer' (Заказчик), 
                        'customer_post' (должность заказчика), 
                        'customer_fio' (ФИО заказчика), 
                        'company' (Исполнитель), 
                        'director_post' (должность руководителя исполнителя), 
                        'director' (ФИО руководителя исполнителя).
                        Текст: {t_context}"""
                    }],
                    response_format={ "type": "json_object" }
                )
                st.session_state.t_info = json.loads(res.choices[0].message.content)
        else: st.error("Нет данных!")

    # --- ПРЕВЬЮ ТИТУЛЬНИКА (Редактируемое) ---
    if "t_info" in st.session_state:
        st.info("Проверьте данные:")
        ti = st.session_state.t_info
        ti['contract_no'] = st.text_input("№", ti.get('contract_no'))
        ti['ikz'] = st.text_input("ИКЗ", ti.get('ikz'))
        ti['customer_fio'] = st.text_input("ФИО Заказчика", ti.get('customer_fio'))
        # Кнопка скачивания только титульника
        doc_t = build_title_page(ti)
        buf_t = io.BytesIO(); doc_t.save(buf_t)
        st.download_button("📥 Скачать Титульник", buf_t.getvalue(), "Title.docx", use_container_width=True)
        
# КОЛОНКА 2: ОТЧЕТ
with col2:
    st.header("📝 2. Отчет (ТЗ)")
    tz_tab1, tz_tab2 = st.tabs(["📁 Файл", "⌨️ Текст"])
    
    with tz_tab1:
        # Добавляем ключ, чтобы файл тоже можно было сбросить
        f_tz = st.file_uploader("Техзадание (DOCX)", type="docx", key=f"u_tz_{st.session_state.reset_counter}")
    
    with tz_tab2:
        # Используем reset_counter для очистки при нажатии кнопки Сброс
        tz_area_key = f"tz_area_{st.session_state.reset_counter}"
        m_tz_area = st.text_area(
            "Текст техзадания:", 
            value=st.session_state.get(tz_area_key, ""), 
            height=150, 
            key=tz_area_key
        )
    
    if st.button("⚙️ Сгенерировать текст", use_container_width=True):
        # Берем текст из окна, если пусто - из файла
        tz_content = m_tz_area.strip() if m_tz_area.strip() else ""
        if not tz_content and f_tz:
            tz_content = get_text_from_file(f_tz)
            
        if tz_content:
            st.session_state.raw_tz_source = tz_content  # СОХРАНЯЕМ ДЛЯ ПОШАГОВОЙ СБОРКИ
            seg_res = client.chat.completions.create(
                model=GEMINI_MODEL,
                messages=[{"role": "user", "content": f"Раздели на блоки с тегом [END_BLOCK]: {tz_content}"}]
            )
            steps = [s.strip() for s in seg_res.choices[0].message.content.split('[END_BLOCK]') if s.strip()]

            instruction = """Роль и контекст:
                    Ты — ассистент юриста по договорной работе. Перед тобой текст Технического Задания (ТЗ), который был написан в будущем времени как план работ. Сейчас его нужно превратить в черновик отчета о выполнении. Твоя задача — действовать как автомат по замене времени и удалению лишних слов, не меняя структуру и терминологию документа. Объём документа может быть очень большим (до 20 страниц) — это нормально, ты должен сохранить весь текст полностью, ничего не выбрасывая и не пересказывая.
                    
                    Инструкция (Что нужно сделать):
                    Необходимо полностью переработать текст ТЗ в текст отчета, выполненного в прошедшем времени, со следующими важными исключениями.
                    
                    Правила обработки (Набор правил):
                    
                    Неприкосновенность заголовков (Важно!): Все заголовки пунктов и подпунктов ТЗ должны остаться в настоящем времени (как в оригинале). Их менять нельзя.
                    
                    Пример: Заголовок «Предоставление транспортных услуг...» должен остаться без изменений.
                    
                    Применяй это правило ко всем уровням заголовков.
                    
                    Основное время (Тело пунктов): Весь описательный текст, следующий за заголовком (внутри пункта), нужно переписать в прошедшее время.
                    
                    Пример: «Исполнитель организует доставку...» -> «Исполнитель организовал доставку...».
                    
                    Чистка текста (Удаление модального мусора): В отчете не должно быть слов, указывающих на долженствование или ограничения из ТЗ. Их нужно удалять или заменять, не искажая сути:
                    
                    Слова для удаления: должен, обязан, нужно, необходимо, следует.
                    
                    Пример: «Исполнитель обязан предоставить отчет» -> «Исполнитель предоставил отчет».
                    
                    Слова для удаления (если они не влияют на цифры): более, менее, не более, не менее, свыше (часто они просто указывают на план, в отчете важны конкретные цифры).
                    
                    Пример: «Поставлено не менее 10 ящиков» -> «Поставлено 10 ящиков» (если факт совпадает с минимумом; если поставлено больше, лучше сохранить факт: «Поставлено 12 ящиков»).
                    
                    Неизменность данных и объёма: Все, что не является глаголами или мусорными словами из п.3, должно остаться нетронутым:
                    
                    Сроки (числа), адреса, имена, названия организаций, специфические термины, номенклатурные номера — все остается как в оригинале ТЗ.
                    
                    Важно: ни в коем случае не сокращай текст, не убирай предложения, не пересказывай своими словами. Сохраняй исходный объём и все детали, даже если текст очень длинный. Просто заменяй времена и удаляй указанные слова.
                    
                    Работа с описаниями процессов: Длинные описания того, как надо делать, превращаются в описание того, как было сделано.
                    
                    Формат вывода:
                    Выведи полностью переработанный текст. Начинай с первого заголовка документа. Не добавляй никаких вступлений, комментариев или пояснений в квадратных скобках. 
                    Только чистый текст отчета."""
            
            # Вызываем Gemini с твоей инструкцией и текстом ТЗ
            res = client.chat.completions.create(
                model=GEMINI_MODEL,
                messages=[{"role": "user", "content": f"{instruction}\n\n{tz_content}"}]
            )
            st.session_state.raw_report_body = res.choices[0].message.content
            
            # Сохраняем результат (у Gemini это res.text)
            st.session_state.raw_report_body = res.choices[0].message.content
            
            final_text_parts = []
            pb = st.progress(0)
            status_text = st.empty()
            
            for i, step in enumerate(steps):
                status_text.text(f"Обработка блока {i+1} из {len(steps)}...")
                part = smart_generate_step_strict(step, st.session_state.get('raw_requirements', ''))
                final_text_parts.append(part)
                pb.progress((i + 1) / len(steps))
            
            st.session_state.raw_report_body = res.choices[0].message.content
        else:
            st.warning("Данные ТЗ отсутствуют")

    if "raw_report_body" in st.session_state:
        st.session_state.raw_report_body = st.text_area("Черновик:", st.session_state.raw_report_body, height=300)

# КОЛОНКА 3: ТРЕБОВАНИЯ
with col3:
    st.header("📋 3. Требования")
    if st.button("🔍 Выделить требования", use_container_width=True):
        if "raw_tz_source" in st.session_state:
            res = client.chat.completions.create(
                model=GEMINI_MODEL,
                messages=[{"role": "user", "content": f"Выпиши требования к документам: {st.session_state.raw_tz_source}"}]
            )
            st.session_state.raw_requirements = res.choices[0].message.content

    if "raw_requirements" in st.session_state:
        st.session_state.raw_requirements = st.text_area("Требования:", st.session_state.raw_requirements, height=300)

# НИЖНИЙ БЛОК: СБОРКА
st.divider()
f_col1, f_col2 = st.columns(2)

with f_col1:
    if st.button("🚀 СОБРАТЬ ПОЛНЫЙ ОТЧЕТ (КАК ЕСТЬ)", use_container_width=True):
        if "t_info" in st.session_state:
            doc = create_final_report(st.session_state.t_info, st.session_state.get('raw_report_body', ''), st.session_state.get('raw_requirements', ''))
            buf = io.BytesIO(); doc.save(buf)
            st.session_state.full_file = buf.getvalue()

with f_col2:
    if st.button("🚀 ЗАПУСТИТЬ ПОШАГОВУЮ СБОРКУ", use_container_width=True):
        if "t_info" in st.session_state and st.session_state.get('raw_tz_source'):
            
            # Разрезаем по пунктам типа 1.1., 2.1.
            steps = [s.strip() for s in re.split(r'\n(?=\d+\.\d+)', st.session_state.raw_tz_source) if s.strip()]
            
            final_text_parts = []
            pb = st.progress(0)
            for i, step in enumerate(steps):
                part = smart_generate_step_strict(step, st.session_state.get('raw_requirements', ''))
                final_text_parts.append(part)
                pb.progress((i + 1) / len(steps))
            
            # Соединяем один раз
            full_smart_text = "\n\n".join(final_text_parts)
            doc = create_final_report(st.session_state.t_info, full_smart_text, st.session_state.get('raw_requirements', ''))
            buf = io.BytesIO()
            doc.save(buf)
            st.session_state.smart_file = buf.getvalue()
            st.success("Умная сборка завершена!")
if "full_file" in st.session_state:
    st.download_button("📥 Скачать обычный", st.session_state.full_file, "Report.docx")
if "smart_file" in st.session_state:
    st.download_button("📥 СКАЧАТЬ УМНЫЙ ОТЧЕТ", st.session_state.smart_file, "Smart_Report.docx")
















